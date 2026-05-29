#!/usr/bin/env python3
"""Regenerate 测试文件.docx with test environment / resources."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "测试文件.docx"


def set_run_font(run, name: str = "微软雅黑", size_pt: int = 11) -> None:
    run.font.name = name
    run.font.size = Pt(size_pt)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_run_font(run, size_pt=16 if level == 1 else 14)


def add_para(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    set_run_font(run)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        set_run_font(run)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            for run in p.runs:
                run.bold = True
                set_run_font(run, size_pt=10)
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            cell = table.rows[ri].cells[ci]
            cell.text = val
            for p in cell.paragraphs:
                for run in p.runs:
                    set_run_font(run, size_pt=10)


def build_document() -> Document:
    doc = Document()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("数据恢复软件 — 测试准备说明")
    tr.bold = True
    set_run_font(tr, size_pt=18)

    add_para(doc, "本文档说明测试环境与测试资源。")

    add_heading(doc, "1. 测试环境要求", level=1)
    add_table(
        doc,
        ["项目", "要求"],
        [
            ["操作系统", "Win10、Win11（64 位）为主"],
            ["32 位系统", "使用优先级低（Win11 不支持 32 位）"],
            ["Win11", "支持"],
            ["Win10", "占比较高，优先覆盖"],
            ["Mac", "本期不在测试范围"],
            ["运行方式", "恢复软件建议以管理员身份运行"],
        ],
    )
    add_para(doc, "当前主力测试机：Win10 台式机。", bold=True)

    add_heading(doc, "2. 测试资源（设备与磁盘）", level=1)
    add_para(doc, "2.1 本机台式（当前执行环境）", bold=True)
    add_para(doc, "Win10 台式\nSSD   240GB   C盘 + D盘\nHDD   500GB   E盘 + F盘 + G盘(5GB) + H盘(10GB)\nV盘   约 8GB   虚拟盘 VHDX\n外接   I盘(EAGET 465GB) + J盘(U盘 232GB) + K盘(SD存储卡 16GB)")
    add_table(
        doc,
        ["设备/磁盘", "类型", "容量", "盘符", "用途", "执行方"],
        [
            ["Win10 台式 · SSD", "固态硬盘", "240GB", "C:、D:", "系统、数据/导出", "不测"],
            ["Win10 台式 · HDD", "机械硬盘", "约 500GB", "E:、F:", "业务数据", "不测"],
            ["Win10 台式 · HDD", "机械硬盘", "5GB 分区", "G: (Test5GB)", "小容量功能测试", "后端自测（P0）"],
            ["Win10 台式 · HDD", "机械硬盘", "10GB 分区", "H: (Test10GB)", "备用测试分区", "测试人员"],
            ["Win10 台式 · VHDX", "虚拟盘", "约 8GB", "V: (CursorVHD)", "虚拟盘场景", "测试人员"],
            ["外接 · EAGET", "外接硬盘", "约 465GB", "I:", "外接移动硬盘", "测试人员"],
            ["外接 · U 盘", "外接闪存盘（U 盘）", "约 232GB", "J:", "外接 U 盘", "测试人员"],
            ["外接 · 闪迪 SD 卡", "SD 存储卡（16GB）", "约 14.8GB 可用", "K:", "SD 卡 + 读卡器", "测试人员"],
            ["笔记本 A / B", "—", "—", "—", "其它 Win10/Win11 环境", "后续扩展"],
        ],
    )
    add_para(
        doc,
        "说明：G、E、F、H 在同一块约 500GB 机械硬盘上；后端只扫 G: 分区，勿选整盘。"
        "V / I / J / K 由测试人员在对应介质上执行。",
    )
    add_para(doc, "V: 虚拟盘文件：D:\\virtual-disks\\cursor_test_vdisk.vhdx")

    add_heading(doc, "3. 当前扫描范围（硬盘扫描专项）", level=1)
    add_table(
        doc,
        ["盘符", "存储类型", "容量", "执行方"],
        [
            ["G:", "机械硬盘（HDD 测试分区）", "5GB", "后端自测（P0）"],
            ["V:", "虚拟盘（VHDX）", "约 8GB", "测试人员"],
            ["H:", "机械硬盘（HDD 测试分区）", "10GB", "测试人员"],
            ["I:", "外接硬盘（EAGET）", "约 465GB", "测试人员"],
            ["J:", "外接 U 盘", "约 232GB", "测试人员"],
            ["K:", "SD 存储卡（闪迪 16GB + 读卡器）", "约 14.8GB 可用", "测试人员"],
            ["C: / D:", "固态硬盘（SSD）", "100GB + 约 124GB", "不测"],
            ["E: / F:", "机械硬盘（HDD 业务分区）", "183GB + 约 233GB", "不测"],
        ],
    )

    add_para(doc, "分工：G: 最稳定，用例 xlsx 仅给后端自测；虚拟盘及外接盘由测试人员执行。", bold=True)

    add_heading(doc, "4. 测试场景类型", level=1)
    add_para(
        doc,
        "产品 6.0 要求：回收站删除、隐藏文件、虚拟盘内已删除文件应在扫描结果「所有文件」中正确展示。"
        "G: 一轮测试（32 条）后，对场景模型做收敛：保留差异维度，合并重复验收项。",
    )

    add_heading(doc, "4.1 原模型（5 情况 × 4 状态）【已收敛，仅作对照】", level=2)
    add_para(
        doc,
        "xlsx「5情况×4状态」与「一轮测试（G）」sheet 仍保留历史 ID，新轮次优先采用 §4.3 收敛模型。",
        bold=True,
    )
    add_table(
        doc,
        ["序号", "场景", "操作要点", "收敛说明"],
        [
            ["1", "普通删除 → 回收站保留", "Delete，不清空回收站", "合并为 S1；隐藏样本并入 S1 抽检"],
            ["2", "删除 → 清空回收站", "Delete 后立即清空回收站", "合并为 S2；隐藏样本并入 S2 抽检"],
            ["3", "隐藏文件（未删除）", "设 Hidden，不删除", "合并为 H1（1 条行为确认）"],
            ["4", "隐藏 + 回收站保留", "Hidden + Delete", "不再单独 4 状态；并入 S1"],
            ["5", "隐藏 + 清空回收站", "Hidden + Delete + 清空", "不再单独 4 状态；并入 S2"],
            ["6", "嵌套文件夹路径", "Delete→回收站保留；验 1/3/5/10 层路径", "收敛为 P1：仅 L01 + L10"],
            ["可选", "永久删除", "Shift+Delete", "合并为 S3；与 006 隐藏永久删除合并抽检"],
        ],
    )

    add_heading(doc, "4.2 一轮测试结论（G: Test5GB）", level=2)
    add_table(
        doc,
        ["维度", "结果", "说明"],
        [
            ["执行规模", "32 条，通过 20 / 失败 12（约 62.5%）", "见 xlsx「一轮测试（G）」"],
            ["检出（-1）", "回收站类基本通过", "001/002/004/005 的检出多为 P"],
            ["原路径 / 预览 / 完整性", "回收站类集中失败", "-2/-3/-4 在 001/002/004/005 上多为 F"],
            ["永久删除", "全过", "DISK_PERM_003、DISK_HIDDEN_006 各 4/4"],
            ["隐藏未删", "全过", "DISK_HIDDEN_003 符合「不进丢失列表」"],
            ["嵌套路径", "全过", "L01/L03/L05/L10 路径展示正常"],
            ["研发优先级", "P0：回收站链路路径与预览", "路径修好后各抽 1 个 docx 做恢复完整性"],
        ],
    )
    add_para(
        doc,
        "注意：样本在 G:\\图片、G:\\文档、G:\\演示 等分类目录下，验收路径应写完整路径（如 G:\\文档\\G_docx_001_*.docx），"
        "勿再按「G:\\根目录\\文件名」填写。",
    )

    add_heading(doc, "4.3 收敛后推荐场景（下一轮起执行）", level=2)
    add_para(doc, "约 8～10 条核心用例，覆盖原 32 条的主要风险面。", bold=True)
    add_table(
        doc,
        ["新 ID", "场景", "操作", "验收（合并后）", "代表样本"],
        [
            ["S1", "回收站保留", "Delete，不清空回收站", "① 能检出 ② 原路径正确 ③ docx 可预览", "各分类目录 1 个 docx；可选 1 个隐藏文件"],
            ["S2", "回收站清空", "Delete 后清空回收站", "同 S1", "同 S1"],
            ["S3", "永久删除", "Shift+Delete", "同 S1（本轮已全过，可降频抽检）", "1 个 docx + 可选 1 个隐藏"],
            ["H1", "隐藏未删", "Hidden，不删除", "扫描结果中不应作为丢失项展示", "任意 1 个已设隐藏的文件"],
            ["P1a", "嵌套路径（浅）", "Delete→回收站保留", "路径含 G:\\nest_L1\\…", "G_nest_L01_*.pptx"],
            ["P1b", "嵌套路径（深）", "Delete→回收站保留", "路径含完整 10 层 nest_L1～L10", "G_nest_L10_*.pptx"],
            ["R1", "恢复完整性（抽测）", "在 S1/S2/S3 各选 1 个 docx", "导出/恢复后本地可打开、内容完整", "不单独 ×4 状态"],
            ["F1", "多格式抽测（可选）", "在 S1 下增测", "pptx、jpg、pdf 各 1 个可预览", "对齐产品文档类型筛选"],
        ],
    )
    add_para(doc, "文档类型覆盖（产品筛选）：PDF/DOC/DOCX/XLS/XLSX/PPT/PPTX、TXT/RTF/ODT/PAGES/NUMBERS/KEY；样本脚本见 stage_user_preview_resources.ps1。")

    add_heading(doc, "5. 测试原则", level=1)
    add_bullets(
        doc,
        [
            "后端自测仅在 G: 进行破坏性操作。",
            "C/D/E/F 禁止批量删除测试样本。",
            "样本按格式放在 G:\\图片、G:\\文档、G:\\表格、G:\\演示、G:\\文本 等分类文件夹（见后端_G盘测试资源准备.md）。",
            "收敛模型下：每条场景合并「检出 + 路径 + 预览」为一次验收；完整性单独抽测（R1）。",
        ],
    )

    add_heading(doc, "6. 验收标准（摘要）", level=1)
    add_bullets(
        doc,
        [
            "P0：S1/S2 路径与预览通过；H1 行为符合产品定义。",
            "不允许出现崩溃、扫描中断、结果页长期空白等阻断问题。",
            "删除文件应能检出；原路径应为删除前完整路径（含子目录）；Open XML（docx/xlsx/pptx）优先验证预览。",
        ],
    )

    return doc


def main() -> None:
    doc = build_document()
    candidates = [OUTPUT, OUTPUT.with_name("测试文件_精简版.docx")]
    last_err: Exception | None = None
    for path in candidates:
        try:
            doc.save(path)
            print(f"Wrote {path}")
            if path != OUTPUT:
                print("原文件被占用，请先关闭 Word 后手动替换，或将完善版重命名为 测试文件.docx")
            return
        except PermissionError as e:
            last_err = e
    raise SystemExit(f"无法写入 docx: {last_err}")


if __name__ == "__main__":
    main()
